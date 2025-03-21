"""
Document loader module for handling various document types.
"""
import os
import hashlib
import time
from pathlib import Path
import fitz  # PyMuPDF
import docx
import pandas as pd
import pytesseract
from PIL import Image
import io
import logging
import json
import concurrent.futures
from multiprocessing import Lock
from typing import Dict, Any, List, Optional, Tuple
# Change this:
import fitz  # PyMuPDF

# To this:

from config.config import RAW_DATA_DIR, PROCESSED_DATA_DIR, OCR_CACHE_DIR, OCR_PARALLEL_JOBS
from src.utils.logger import setup_logger
from src.utils.resource_monitor import log_memory_usage

logger = setup_logger(__name__, "document_loader.log")

# Create a lock for thread-safe access to the OCR cache
ocr_cache_lock = Lock()

class DocumentLoader:
    """Document loader that handles various file formats."""
    
    def __init__(self, raw_dir: Optional[str] = None, processed_dir: Optional[str] = None, 
                ocr_cache_dir: Optional[str] = None):
        """
        Initialize document loader.
        
        Args:
            raw_dir: Directory for raw documents. Defaults to config.
            processed_dir: Directory for processed documents. Defaults to config.
            ocr_cache_dir: Directory for OCR cache. Defaults to config.
        """
        self.raw_dir = Path(raw_dir) if raw_dir else RAW_DATA_DIR
        self.processed_dir = Path(processed_dir) if processed_dir else PROCESSED_DATA_DIR
        self.ocr_cache_dir = Path(ocr_cache_dir) if ocr_cache_dir else OCR_CACHE_DIR
        
        # Number of parallel OCR jobs (leave 2 cores free)
        self.ocr_parallel_jobs = OCR_PARALLEL_JOBS
        
        # Ensure directories exist
        os.makedirs(self.raw_dir, exist_ok=True)
        os.makedirs(self.processed_dir, exist_ok=True)
        os.makedirs(self.ocr_cache_dir, exist_ok=True)
        
        logger.info(f"Document loader initialized with raw_dir={self.raw_dir}, "
                   f"processed_dir={self.processed_dir}, "
                   f"ocr_cache_dir={self.ocr_cache_dir}, "
                   f"ocr_parallel_jobs={self.ocr_parallel_jobs}")
    
    def _get_cache_key(self, image_data: bytes) -> str:
        """
        Generate a cache key from image data.
        
        Args:
            image_data: Image data
            
        Returns:
            Cache key
        """
        return hashlib.md5(image_data).hexdigest()
    
    def _get_cached_ocr(self, cache_key: str) -> Optional[str]:
        """
        Get OCR result from cache.
        
        Args:
            cache_key: Cache key
            
        Returns:
            OCR text if cached, None otherwise
        """
        cache_file = self.ocr_cache_dir / f"{cache_key}.json"
        if cache_file.exists():
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cached_data = json.load(f)
                    logger.info(f"OCR cache hit: {cache_key}")
                    return cached_data.get('text', '')
            except Exception as e:
                logger.warning(f"Error reading OCR cache: {e}")
        
        return None
    
    def _save_to_ocr_cache(self, cache_key: str, text: str) -> None:
        """
        Save OCR result to cache.
        
        Args:
            cache_key: Cache key
            text: OCR text
        """
        cache_file = self.ocr_cache_dir / f"{cache_key}.json"
        try:
            with ocr_cache_lock:
                with open(cache_file, 'w', encoding='utf-8') as f:
                    json.dump({'text': text}, f)
            logger.info(f"Saved to OCR cache: {cache_key}")
        except Exception as e:
            logger.warning(f"Error saving to OCR cache: {e}")
    
    def purge_ocr_cache(self) -> Tuple[bool, int]:
        """
        Purge all cached OCR results.
        
        Returns:
            tuple: (success_status, count_of_files_removed)
        """
        try:
            # Check if OCR cache directory exists
            if not os.path.exists(self.ocr_cache_dir):
                logger.info(f"OCR cache directory does not exist: {self.ocr_cache_dir}")
                return True, 0
            
            # Count and delete all cache files
            file_count = 0
            for filename in os.listdir(self.ocr_cache_dir):
                if filename.endswith('.json'):
                    filepath = os.path.join(self.ocr_cache_dir, filename)
                    os.remove(filepath)
                    file_count += 1
            
            logger.info(f"Purged {file_count} files from OCR cache")
            return True, file_count
            
        except Exception as e:
            logger.error(f"Error purging OCR cache: {e}")
            return False, 0
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document from the given file path.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Document data including metadata and text content
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Loading document: {file_path}")
        
        # Extract file extension and call appropriate handler
        file_extension = file_path.suffix.lower()
        
        handlers = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,  # May need additional handling for older .doc files
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.csv': self._process_csv,
            '.txt': self._process_text,
        }
        
        if file_extension in handlers:
            result = handlers[file_extension](file_path)
            log_memory_usage(logger)
            return result
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """
        Process a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Document data including metadata and text content
        """
        logger.info(f"Processing PDF file: {file_path}")
        
        doc_data = {
            'document_id': str(file_path.stem),
            'file_name': file_path.name,
            'file_type': 'pdf',
            'content': [],
            'metadata': {},
            'images': []
        }
        
        try:
            # Open the PDF
            pdf_document = fitz.open(file_path)
            
            # Extract metadata
            doc_data['metadata'] = {
                'title': pdf_document.metadata.get('title', ''),
                'author': pdf_document.metadata.get('author', ''),
                'creation_date': pdf_document.metadata.get('creationDate', ''),
                'modification_date': pdf_document.metadata.get('modDate', ''),
                'page_count': len(pdf_document)
            }
            
            # Create a unified approach to ensure we get all text from a PDF
            # First, extract the text normally from all pages
            page_contents = []
            for page_num, page in enumerate(pdf_document):
                page_text = page.get_text()
                page_contents.append({
                    'page_num': page_num + 1,
                    'text': page_text,
                    'page': page,
                    'needs_ocr': len(page_text.strip()) < 100  # Flag pages with little text
                })
            
            # Collect all pages that need OCR and all images that need OCR
            ocr_tasks = []
            
            # Pages that need OCR
            for page_data in page_contents:
                if page_data['needs_ocr']:
                    ocr_tasks.append(('page', page_data['page'], page_data['page_num']))
            
            # Images from the PDF
            for page_num, page in enumerate(pdf_document):
                image_list = page.get_images(full=True)
                
                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Only process images of reasonable size (to avoid tiny icons)
                    img = Image.open(io.BytesIO(image_bytes))
                    width, height = img.size
                    if width > 100 and height > 100:  # Skip very small images
                        ocr_tasks.append(('image', image_bytes, (page_num + 1, img_index)))
            
            logger.info(f"Collected {len(ocr_tasks)} OCR tasks")
            
            # Process OCR tasks in parallel
            ocr_results = {}
            if ocr_tasks:
                with concurrent.futures.ThreadPoolExecutor(max_workers=self.ocr_parallel_jobs) as executor:
                    future_to_task = {
                        executor.submit(self._perform_ocr_task, task_type, content, identifier): 
                        (task_type, identifier) 
                        for task_type, content, identifier in ocr_tasks
                    }
                    
                    for future in concurrent.futures.as_completed(future_to_task):
                        task_type, identifier, ocr_text = future.result()
                        if ocr_text:
                            ocr_results[(task_type, identifier)] = ocr_text
            
            # Update page contents with OCR results
            for page_data in page_contents:
                page_num = page_data['page_num']
                if page_data['needs_ocr'] and ('page', page_data['page_num']) in ocr_results:
                    page_data['text'] = ocr_results[('page', page_data['page_num'])]
                
                # Add image OCR text if available
                for (task_type, identifier), ocr_text in ocr_results.items():
                    if task_type == 'image' and identifier[0] == page_num:
                        img_index = identifier[1]
                        
                        # Add to images list
                        doc_data['images'].append({
                            'page': page_num - 1,  # 0-indexed
                            'index': img_index,
                            'ocr_text': ocr_text
                        })
                        
                        # Augment page text with image OCR
                        page_data['text'] += f"\n\n[Image OCR Text: {ocr_text}]"
            
            # Create final content
            for page_data in page_contents:
                doc_data['content'].append({
                    'page_num': page_data['page_num'],
                    'text': page_data['text']
                })
            
            # Close the document
            pdf_document.close()
            
            logger.info(f"PDF processing complete: {file_path} - {len(doc_data['content'])} pages")
            return doc_data
            
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {e}")
            raise
            
    def _perform_ocr_task(self, task_type: str, content: Any, identifier: Any) -> Tuple[str, Any, str]:
        """
        Perform OCR on a task.
        
        Args:
            task_type: 'page' or 'image'
            content: Page object or image bytes
            identifier: Page number or (page_num, img_index) tuple
            
        Returns:
            tuple: (task_type, identifier, ocr_text)
        """
        try:
            if task_type == 'page':
                page = content
                page_num = identifier
                
                # Convert page to image
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))  # Lower resolution
                img_bytes = pix.tobytes()
                
                # Try cache first
                cache_key = self._get_cache_key(img_bytes)
                cached_text = self._get_cached_ocr(cache_key)
                
                if cached_text is not None:
                    return (task_type, page_num, cached_text)
                
                # Perform OCR
                image = Image.open(io.BytesIO(img_bytes))
                ocr_text = pytesseract.image_to_string(image)
                image.close()
                
                # Cache the result
                self._save_to_ocr_cache(cache_key, ocr_text)
                
                return (task_type, page_num, ocr_text)
                
            elif task_type == 'image':
                image_bytes = content
                page_img_index = identifier
                
                # Try cache first
                cache_key = self._get_cache_key(image_bytes)
                cached_text = self._get_cached_ocr(cache_key)
                
                if cached_text is not None:
                    return (task_type, page_img_index, cached_text)
                
                # Perform OCR
                image = Image.open(io.BytesIO(image_bytes))
                ocr_text = pytesseract.image_to_string(image)
                image.close()
                
                # Only store if it found text
                if len(ocr_text.strip()) > 10:
                    # Cache the result
                    self._save_to_ocr_cache(cache_key, ocr_text)
                    return (task_type, page_img_index, ocr_text)
                
                return (task_type, page_img_index, '')
            
            return (task_type, identifier, '')
            
        except Exception as e:
            logger.error(f"OCR error for {task_type} {identifier}: {e}")
            return (task_type, identifier, '')
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """
        Process a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Document data including metadata and text content
        """
        logger.info(f"Processing DOCX file: {file_path}")
        
        doc_data = {
            'document_id': str(file_path.stem),
            'file_name': file_path.name,
            'file_type': 'docx',
            'content': [],
            'metadata': {}
        }
        
        try:
            # Open the document
            document = docx.Document(file_path)
            
            # Extract metadata
            core_properties = document.core_properties
            doc_data['metadata'] = {
                'title': core_properties.title if hasattr(core_properties, 'title') else '',
                'author': core_properties.author if hasattr(core_properties, 'author') else '',
                'created': str(core_properties.created) if hasattr(core_properties, 'created') else '',
                'modified': str(core_properties.modified) if hasattr(core_properties, 'modified') else '',
                'paragraph_count': len(document.paragraphs),
                'table_count': len(document.tables)
            }
            
            # Extract content from paragraphs
            full_text = []
            for para in document.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract content from tables
            table_texts = []
            for table in document.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                table_texts.append("\n".join(table_text))
            
            # Combine all text
            doc_data['content'].append({
                'text': "\n\n".join(full_text + ["\n\nTABLES:\n"] + table_texts if table_texts else full_text)
            })
            
            logger.info(f"DOCX processing complete: {file_path}")
            return doc_data
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {e}")
            raise
    
    def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """
        Process an Excel file.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Document data including metadata and text content
        """
        logger.info(f"Processing Excel file: {file_path}")
        
        doc_data = {
            'document_id': str(file_path.stem),
            'file_name': file_path.name,
            'file_type': 'excel',
            'content': [],
            'metadata': {},
            'sheets': []
        }
        
        try:
            # Read Excel file with pandas
            excel_file = pd.ExcelFile(file_path)
            
            # Extract metadata
            doc_data['metadata'] = {
                'sheet_names': excel_file.sheet_names,
                'sheet_count': len(excel_file.sheet_names)
            }
            
            # Process each sheet
            all_text_content = []
            
            for sheet_name in excel_file.sheet_names:
                logger.info(f"Processing sheet: {sheet_name}")
                
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Convert sheet to structured text
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                sheet_text += f"Column headers: {', '.join(str(col) for col in df.columns)}\n\n"
                
                # Add sample data (first 10 rows)
                sample_rows = min(10, len(df))
                if sample_rows > 0:
                    sheet_text += "Sample data:\n"
                    sample_df = df.head(sample_rows)
                    sheet_text += sample_df.to_string() + "\n\n"
                
                # Save the full sheet data for potential further processing
                doc_data['sheets'].append({
                    'name': sheet_name,
                    'dataframe': df.to_dict(orient='records')
                })
                
                all_text_content.append(sheet_text)
            
            # Combine all sheets into one content item
            doc_data['content'].append({
                'text': "\n\n".join(all_text_content)
            })
            
            logger.info(f"Excel processing complete: {file_path} - {len(doc_data['sheets'])} sheets")
            return doc_data
            
        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {e}")
            raise
    
    def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """
        Process a CSV file.
        
        Args:
            file_path: Path to the CSV file
            
        Returns:
            Document data including metadata and text content
        """
        logger.info(f"Processing CSV file: {file_path}")
        
        doc_data = {
            'document_id': str(file_path.stem),
            'file_name': file_path.name,
            'file_type': 'csv',
            'content': [],
            'metadata': {}
        }
        
        try:
            # Read CSV file
            df = pd.read_csv(file_path)
            
            # Extract metadata
            doc_data['metadata'] = {
                'row_count': len(df),
                'column_count': len(df.columns),
                'columns': list(df.columns)
            }
            
            # Convert to structured text
            text_content = f"CSV File: {file_path.name}\n"
            text_content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
            text_content += f"Column headers: {', '.join(str(col) for col in df.columns)}\n\n"
            
            # Add sample data (first 10 rows)
            sample_rows = min(10, len(df))
            if sample_rows > 0:
                text_content += "Sample data:\n"
                sample_df = df.head(sample_rows)
                text_content += sample_df.to_string() + "\n\n"
            
            # Store the full dataframe for potential further processing
            doc_data['dataframe'] = df.to_dict(orient='records')
            
            # Add content for text extraction
            doc_data['content'].append({
                'text': text_content
            })
            
            logger.info(f"CSV processing complete: {file_path} - {doc_data['metadata']['row_count']} rows")
            return doc_data
            
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {e}")
            raise
    
    def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """
        Process a plain text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Document data including metadata and text content
        """
        logger.info(f"Processing text file: {file_path}")
        
        doc_data = {
            'document_id': str(file_path.stem),
            'file_name': file_path.name,
            'file_type': 'text',
            'content': [],
            'metadata': {}
        }
        
        try:
            # Read text file
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                text_content = f.read()
            
            # Extract basic metadata
            doc_data['metadata'] = {
                'file_size': os.path.getsize(file_path),
                'line_count': text_content.count('\n') + 1,
                'character_count': len(text_content)
            }
            
            # Add content
            doc_data['content'].append({
                'text': text_content
            })
            
            logger.info(f"Text file processing complete: {file_path} - {doc_data['metadata']['character_count']} characters")
            return doc_data
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise